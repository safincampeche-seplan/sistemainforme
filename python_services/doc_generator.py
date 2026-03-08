from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from models import ExportRequest, NarrativeItem
import io
import re

def sanitize_xml(text: str) -> str:
    if not text:
        return ""
    # Romoves NULL bytes and control chars outside ASCII/XML compatible sets
    _illegal_xml_chars_RE = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1F\uD800-\uDFFF\uFFFE\uFFFF]')
    return _illegal_xml_chars_RE.sub('', text)

def generate_word_doc(request: ExportRequest) -> io.BytesIO:
    doc = Document()
    
    # Estilos por defecto (Simulando Averta con Calibri/Inter si no está instalada, 
    # pero configurando el nombre para que Word lo reconozca)
    
    current_title_id = None
    current_theme_id = None
    current_subtheme_id = None
    current_program_name = None

    for item in request.items:
        # Título (H1 equivalente)
        if current_title_id != item.title_code:
            current_title_id = item.title_code
            p = doc.add_paragraph()
            run = p.add_run(f"{item.title_code}. {sanitize_xml(item.title_name).upper()}")
            run.bold = True
            run.font.name = 'Averta Bold'
            run.font.size = Pt(21)
            run.font.color.rgb = RGBColor.from_string(request.title_color.replace("#", ""))
            p.space_after = Pt(12)

        # Tema (H2 equivalente)
        if item.theme_name and current_theme_id != item.theme_code:
            current_theme_id = item.theme_code
            p = doc.add_paragraph()
            run = p.add_run(f"{item.title_code}.{item.theme_code}. {sanitize_xml(item.theme_name).upper()}")
            run.font.name = 'Averta Regular'
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor.from_string(request.theme_color.replace("#", ""))

        # Subtema (H3 equivalente)
        if item.subtheme_name and current_subtheme_id != item.subtheme_code:
            current_subtheme_id = item.subtheme_code
            current_program_name = None # Resetear el programa cuando cambie el subtema
            p = doc.add_paragraph()
            run = p.add_run(f"{item.title_code}.{item.theme_code}.{item.subtheme_code}. {sanitize_xml(item.subtheme_name)}")
            run.font.name = 'Averta Regular'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor.from_string(request.subtheme_color.replace("#", ""))

        # Programa Presupuestario (H4 equivalente)
        if item.program_name and current_program_name != item.program_name:
            current_program_name = item.program_name
            p = doc.add_paragraph()
            run = p.add_run(f"{sanitize_xml(item.program_name)}")
            run.font.name = 'Averta Regular'
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor.from_string(request.subtheme_color.replace("#", ""))
            p.space_before = Pt(6)
            p.space_after = Pt(4)

        # Narrativa (Cuerpo)
        if item.content:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
            p.space_before = Pt(8)
            p.space_after = Pt(8)
            
            # Procesar saltos de línea (reales \n o literales \\n si vienen de JSON escapado)
            content = sanitize_xml(item.content).replace("\\n", "\n")
            lines = content.split("\n")
            
            for i, line in enumerate(lines):
                run = p.add_run(line)
                run.font.name = 'Averta Regular'
                run.font.size = Pt(11)
                if i < len(lines) - 1:
                    run.add_break()

        # Destacado (Quote)
        if item.highlighted:
            # Replicar lógica Laravel: Remover todos los saltos de línea y múltiples espacios por un soló espacio
            flattened_highlight = re.sub(r'\s+', ' ', sanitize_xml(item.highlighted)).strip()
            
            # Crear tabla para el destacado (emulando 70% ancho aprox con márgenes fijos)
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.autofit = False
            
            # Asignando 4.5 pulgadas (aprox 70% de la zona imprimible de página letter) a la celda principal
            from docx.shared import Inches
            table.columns[0].width = Inches(4.5)
            
            cell = table.rows[0].cells[0]
            cell.width = Inches(4.5)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = p.add_run(f'"{flattened_highlight}"')
            run.font.name = 'Quatro Slab'
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor.from_string(request.title_color.replace("#", ""))
            
            # Interpolando los 240 twips de Laravel (240/20 = 12 puntos)
            p.space_before = Pt(12)
            p.space_after = Pt(12)

    # Guardar en buffer
    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    return target
